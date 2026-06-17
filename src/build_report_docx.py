"""
build_report_docx.py — render reports/REPORT.md to a Word document reports/REPORT.docx,
embedding the figures from reports/figures/. Mirrors build_report_html.py.

Handles the markdown subset used in REPORT.md: ATX headings (#..###), paragraphs with
inline **bold** / *italic* / `code`, horizontal rules (---), pipe tables (with alignment),
ordered/unordered lists, and image lines (![caption](path)) which are embedded with a caption.
"""

import os
import re

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RP = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "reports")
MD = os.path.join(RP, "REPORT.md")
OUT = os.path.join(RP, "REPORT.docx")

IMG_RE = re.compile(r"^!\[(?P<alt>.*?)\]\((?P<path>.*?)\)\s*$")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def add_inline(paragraph, text):
    """Add text to a paragraph, honouring **bold**, *italic*, and `code` spans."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:|-]+\|?\s*$", line)) and "-" in line


def main():
    with open(MD, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            doc.add_paragraph().add_run().add_break()
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            add_inline(heading, m.group(2))
            i += 1
            continue

        # image
        im = IMG_RE.match(stripped)
        if im:
            path = im.group("path")
            abspath = path if os.path.isabs(path) else os.path.join(RP, path)
            if os.path.exists(abspath):
                doc.add_picture(abspath, width=Inches(5.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            crun = cap.add_run(im.group("alt"))
            crun.italic = True
            crun.font.size = Pt(9)
            crun.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            i += 1
            continue

        # table: a pipe line followed by a separator line
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for j, htext in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(htext)
                run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    val = row[j] if j < len(row) else ""
                    cells[j].paragraphs[0].text = ""
                    add_inline(cells[j].paragraphs[0], val)
            continue

        # unordered list
        lm = re.match(r"^[-*]\s+(.*)$", stripped)
        if lm:
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, lm.group(1))
            i += 1
            continue

        # ordered list
        om = re.match(r"^\d+\.\s+(.*)$", stripped)
        if om:
            para = doc.add_paragraph(style="List Number")
            add_inline(para, om.group(1))
            i += 1
            continue

        # plain paragraph (gather following continuation lines)
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,6}\s|[-*]\s|\d+\.\s|!\[|---)", lines[i].strip()
        ) and not ("|" in lines[i] and i + 1 < n and is_table_sep(lines[i + 1])):
            buf.append(lines[i].strip())
            i += 1
        para = doc.add_paragraph()
        add_inline(para, " ".join(buf))

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
